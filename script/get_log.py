import re
from docx import Document

class get_log:
    def __init__(self,files,month_list,year):
        self.files=files
        self.month_list=month_list
        self.year = year

    def get_log(self):
        list_hostname = []
        list_log = []
        list_severity_value = []
        check_log = 'False'
        count_file = 1
        for file in self.files:
            try:
                print('')
                print('File '+str(count_file)+' of '+str(len(self.files)))
                print('Processing File :')
                print(file)
                read_file = open(file, 'r')
                read_file_list = read_file.readlines()
                
                for file in read_file_list:
                    if re.findall('^hostname (.*)', file):
                        hostname = re.findall('^hostname (.*)', file)
                        hostname = hostname[0]
                        break
                
                log_break = False
                for file in read_file_list:
                    #Check if log has severity below 4
                    for month in self.month_list:
                        #condition has month only
                        if re.findall('^.*('+month+'\s+\d+\s+\d+:.*-[0,1,2,3,4]-.*)', file):
                            log_break = True
                            log = re.findall('^.*('+month+'\s+\d+\s+\d+:.*-[0,1,2,3,4]-.*)', file)
                            log = log[0]
                            severity_value = re.findall('-(\d+)-', file)
                            severity_value = severity_value[0]
                            list_hostname.append(hostname)
                            list_log.append(log)
                            list_severity_value.append(severity_value)
                            check_log = 'True'
                        #condition has both month and year
                        elif re.findall('^.*('+month+'\s+\d+\s+'+self.year+'\s+.*-[0,1,2,3,4]-.*)',file):
                            log_break = True
                            log = re.findall('^.*('+month+'\s+\d+\s+'+self.year+'\s+.*-[0,1,2,3,4]-.*)', file)
                            log = log[0]
                            severity_value = re.findall('-(\d+)-', file)
                            severity_value = severity_value[0]
                            list_hostname.append(hostname)
                            list_log.append(log)
                            list_severity_value.append(severity_value)
                            check_log = 'True'
                        elif re.findall('\S+\s+\d+\s+.*%.*-\d+-',file):
                            log_break = True
                        #break loop
                        elif log_break == True and re.findall('.*#',file):
                            break
                        elif log_break == True and re.findall('^\s*$',file):
                            break
                
                #Check if log has not severity below 4
                if check_log == 'False':
                    list_hostname.append(hostname)
                    list_log.append('-')
                    list_severity_value.append('-')
                
                check_log = 'False'    

            #except NameError:
                #raise
            except:
                pass
            count_file+=1 
        
        print('')
        print('Processing Document')

        write = open('logresult', 'w')
        write.write('Device Name'+' | '+'Log Event'+
                        ' | '+'Severity Value'+' | '+'Severity'+'\n')
        for enum, logtext in enumerate (list_log):
            write.write(list_hostname[enum]+' | '+logtext+
                        ' | '+list_severity_value[enum]+'\n')

        print('')
        print('Saving Document')
        print('Document has been saved to logresult')
'''           
        #using document docx module
        document = Document()
        print('')
        print('Processing Document')
        #add to document
        p = document.add_paragraph('')
        p.add_run('LOG SUMMARY').bold = True
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'Log Event'
        hdr_cells[2].text = 'Severity'
        count_host = 0        
        for hostname in list_hostname:
            row_cells = table.add_row().cells
            row_cells[0].text = hostname
            row_cells[1].text = list_log[count_host]
            count_host+=1

        #save document
        print('Saving Document')
        document.save('show_log.docx')
        print('Document has been saved to show_log.docx')
'''